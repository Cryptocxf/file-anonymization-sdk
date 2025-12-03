"""
Word文件脱敏器
基于python-docx实现Word文档敏感信息脱敏
"""

import logging
from typing import Optional, Dict, Any
from faker import Faker
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.backends import default_backend
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
import base64

from .base_anonymizer import BaseAnonymizer
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities.engine import OperatorConfig

logger = logging.getLogger(__name__)


class WordAnonymizer(BaseAnonymizer):
    """Word文件脱敏器"""
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    SUPPORTED_METHODS = ['fake', 'mask', 'encrypt']
    
    def __init__(self, language: str = 'zh', verbose: bool = False, 
                 output_dir: str = "./anonymized-datas"):
        super().__init__(language, verbose, output_dir)
        
        # 初始化Faker
        if self.language == 'zh':
            self.faker = Faker(locale="zh_CN")
        else:
            self.faker = Faker(locale="en_US")
        
        # 初始化匿名化引擎
        self.anonymizer = AnonymizerEngine()
    
    def anonymize(self, input_path: str, output_path: Optional[str] = None, 
                 method: str = 'mask', encryption_key: Optional[str] = None, 
                 **kwargs) -> str:
        """
        Word文件脱敏
        
        Args:
            input_path: 输入Word路径
            output_path: 输出路径（可选）
            method: 脱敏方法（fake, mask, encrypt）
            encryption_key: 加密密钥（6位数字，仅encrypt方法需要）
        
        Returns:
            str: 输出文件路径
        
        Raises:
            FileValidationError: 文件验证失败
            MethodNotSupportedError: 方法不支持
            EncryptionKeyError: 加密密钥错误
            Exception: 处理错误
        """
        from ..exceptions import (
            FileValidationError, MethodNotSupportedError, EncryptionKeyError
        )
        
        # 验证文件
        is_valid, message = self.validate_file(input_path)
        if not is_valid:
            raise FileValidationError(message)
        
        # 验证方法
        is_valid, message = self.validate_method(method)
        if not is_valid:
            raise MethodNotSupportedError(message)
        
        # 验证加密密钥
        if method == 'encrypt':
            if not encryption_key:
                raise EncryptionKeyError("加密方法需要提供加密密钥")
            if not (len(encryption_key) == 6 and encryption_key.isdigit()):
                raise EncryptionKeyError("加密密钥必须是6位数字")
        
        # 生成输出路径
        if output_path is None:
            output_path = self.get_output_path(input_path, method)
        
        logger.info(f"开始Word脱敏: {input_path} -> {output_path}")
        logger.info(f"脱敏方法: {method}")
        
        try:
            # 根据方法获取操作符配置
            operators = self._get_operators(method, encryption_key)
            
            # 执行脱敏
            total_redactions = self._anonymize_docx(
                input_path, output_path, operators
            )
            
            logger.info(f"Word脱敏完成，共处理 {total_redactions} 处敏感信息")
            return output_path
            
        except Exception as e:
            logger.error(f"Word脱敏失败: {e}")
            raise
    
    def _get_operators(self, method: str, encryption_key: Optional[str] = None) -> Dict[str, OperatorConfig]:
        """获取操作符配置"""
        if method == 'fake':
            return self._get_fake_operators()
        elif method == 'mask':
            return self._get_mask_operators()
        elif method == 'encrypt':
            if not encryption_key:
                raise ValueError("加密方法需要加密密钥")
            return self._get_encrypt_operators(encryption_key)
        else:
            raise ValueError(f"未知的脱敏方法: {method}")
    
    def _get_fake_operators(self) -> Dict[str, OperatorConfig]:
        """伪造数据操作符"""
        return {
            "PERSON": OperatorConfig("custom", {"lambda": lambda x: self.faker.name()}),
            "PHONE_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self.faker.phone_number()}),
            "LOCATION": OperatorConfig("custom", {"lambda": lambda x: f"{self.faker.province()}{self.faker.city()}"}),
            "EMAIL_ADDRESS": OperatorConfig("custom", {"lambda": lambda x: self.faker.safe_email()}),
            "DATE_TIME": OperatorConfig("custom", {"lambda": lambda x: self.faker.past_date().strftime('%Y-%m-%d')}),
            "CREDIT_CARD": OperatorConfig("custom", {"lambda": lambda x: self.faker.credit_card_number()}),
            "US_BANK_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self.faker.credit_card_number()}),
            "DEFAULT": OperatorConfig("replace", {"new_value": "***"}),
        }
    
    def _get_mask_operators(self) -> Dict[str, OperatorConfig]:
        """打码操作符"""
        return {
            "PERSON": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 1)}),
            "PHONE_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 3)}),
            "LOCATION": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 2)}),
            "EMAIL_ADDRESS": OperatorConfig("custom", {"lambda": lambda x: self._mask_email(x)}),
            "DATE_TIME": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 4)}),
            "CREDIT_CARD": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 4)}),
            "US_BANK_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 4)}),
            "DEFAULT": OperatorConfig("custom", {"lambda": lambda x: self._mask_text(x, 2)}),
        }
    
    def _get_encrypt_operators(self, key: str) -> Dict[str, OperatorConfig]:
        """加密操作符"""
        return {
            "PERSON": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "PHONE_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "LOCATION": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "EMAIL_ADDRESS": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "DATE_TIME": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "CREDIT_CARD": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "US_BANK_NUMBER": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
            "DEFAULT": OperatorConfig("custom", {"lambda": lambda x: self._encrypt_text(x, key)}),
        }
    
    def _mask_text(self, text: str, keep_chars: int = 3) -> str:
        """打码处理"""
        if not text or len(text) <= keep_chars:
            return text
        return text[:keep_chars] + '*' * (len(text) - keep_chars)
    
    def _mask_email(self, email: str) -> str:
        """邮箱打码处理"""
        if '@' not in email:
            return self._mask_text(email, 2)
        
        local, domain = email.split('@', 1)
        masked_local = self._mask_text(local, 2)
        return f"{masked_local}@{domain}"
    
    def _encrypt_text(self, text: str, key: str) -> str:
        """加密文本"""
        if not text:
            return text
        
        try:
            # 使用PBKDF2从密钥派生加密密钥
            salt = b'word_anonymizer_salt'
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,
                backend=default_backend()
            )
            key_bytes = key.encode()
            derived_key = kdf.derive(key_bytes)
            
            # 使用AES ECB模式进行确定性加密
            cipher = Cipher(algorithms.AES(derived_key), modes.ECB(), backend=default_backend())
            encryptor = cipher.encryptor()
            
            # 对文本进行填充
            block_size = 16
            padded_text = text.encode()
            padding_length = block_size - (len(padded_text) % block_size)
            padded_text += bytes([padding_length] * padding_length)
            
            # 加密
            encrypted_data = encryptor.update(padded_text) + encryptor.finalize()
            
            return f"ENC:{base64.urlsafe_b64encode(encrypted_data).decode()}"
            
        except Exception as e:
            logger.error(f"加密失败: {e}")
            return text
    
    def _anonymize_docx(self, input_path: str, output_path: str, 
                       operators: Dict[str, OperatorConfig]) -> int:
        """执行DOCX文件脱敏"""
        from docx import Document
        import re
        
        try:
            doc = Document(input_path)
            total_redactions = 0
            
            # 处理段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    original_text = paragraph.text
                    anonymized_text = self._process_text(original_text, operators)
                    
                    if anonymized_text != original_text:
                        total_redactions += self._count_redactions(original_text, anonymized_text)
                        
                        # 保留格式
                        if paragraph.runs:
                            # 清除所有运行
                            for run in paragraph.runs:
                                run.text = ""
                            # 在第一个运行中添加脱敏后的文本
                            paragraph.runs[0].text = anonymized_text
                        else:
                            paragraph.text = anonymized_text
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                original_text = paragraph.text
                                anonymized_text = self._process_text(original_text, operators)
                                
                                if anonymized_text != original_text:
                                    total_redactions += self._count_redactions(original_text, anonymized_text)
                                    
                                    # 保留格式
                                    if paragraph.runs:
                                        for run in paragraph.runs:
                                            run.text = ""
                                        paragraph.runs[0].text = anonymized_text
                                    else:
                                        paragraph.text = anonymized_text
            
            # 保存文档
            doc.save(output_path)
            return total_redactions
            
        except Exception as e:
            logger.error(f"处理DOCX失败: {e}")
            raise
    
    def _process_text(self, text: str, operators: Dict[str, OperatorConfig]) -> str:
        """处理文本"""
        try:
            # 分析敏感信息
            analyzer_results = self.analyzer.analyze(text)
            
            if analyzer_results:
                # 使用匿名化器处理
                anonymized_results = self.anonymizer.anonymize(
                    text=text,
                    analyzer_results=analyzer_results,
                    operators=operators,
                )
                return anonymized_results.text
            else:
                return text
                
        except Exception as e:
            logger.error(f"处理文本失败: {e}")
            return text
    
    def _count_redactions(self, original_text: str, anonymized_text: str) -> int:
        """计算替换次数"""
        # 简单实现：统计改变了多少字符
        if len(original_text) != len(anonymized_text):
            return 1  # 如果长度不同，至少有一次替换
        
        changes = sum(1 for a, b in zip(original_text, anonymized_text) if a != b)
        return 1 if changes > 0 else 0
    
    def extract_text(self, input_path: str) -> str:
        """
        提取Word文本内容
        
        Args:
            input_path: Word文件路径
        
        Returns:
            Word文本内容
        """
        try:
            from docx import Document
            
            doc = Document(input_path)
            text_parts = []
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"提取Word文本失败: {e}")
            return ""